#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script pour générer un rapport Word détaillé du projet FootCamp Dreams
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import json
    from datetime import datetime
except ImportError as e:
    print(f"Erreur d'import: {e}")
    print("Installation des dépendances...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx", "pillow"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import json
    from datetime import datetime

# Créer le document
doc = Document()

# Configuration du style
def set_heading_style(paragraph, level=1, color=None):
    """Applique un style de titre avec couleur optionnelle"""
    if level == 1:
        paragraph.style = 'Heading 1'
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(6)
        if color:
            for run in paragraph.runs:
                run.font.color.rgb = color
    elif level == 2:
        paragraph.style = 'Heading 2'
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(4)
    elif level == 3:
        paragraph.style = 'Heading 3'
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(3)

def add_colored_heading(doc, text, level=1):
    """Ajoute un titre coloré"""
    para = doc.add_paragraph(text)
    set_heading_style(para, level)
    return para

def shade_cell(cell, color):
    """Ajoute une couleur de fond à une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def add_table_with_data(doc, headers, rows, header_color='D3D3D3'):
    """Crée un tableau avec données"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # En-têtes
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        shade_cell(header_cells[i], header_color)
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Données
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)

    return table

# ===== PAGE DE TITRE =====
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('FOOTCAMP DREAMS')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 230, 118)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Application Web de Gestion de Réservations')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(138, 155, 176)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run(f'Rapport Technique\n{datetime.now().strftime("%d %B %Y")}')
run.font.size = Pt(10)

doc.add_paragraph()  # Espace

# ===== TABLE DES MATIÈRES =====
add_colored_heading(doc, '1. PRÉSENTATION GÉNÉRALE', 1)
add_colored_heading(doc, '2. ARCHITECTURE TECHNIQUE', 1)
add_colored_heading(doc, '3. FONCTIONNALITÉS PRINCIPALES', 1)
add_colored_heading(doc, '4. INTERFACE UTILISATEUR', 1)
add_colored_heading(doc, '5. API ET ENDPOINTS', 1)
add_colored_heading(doc, '6. GESTION DES DONNÉES', 1)
add_colored_heading(doc, '7. SÉCURITÉ ET AUTHENTIFICATION', 1)
add_colored_heading(doc, '8. INSTALLATION ET DÉPLOIEMENT', 1)

doc.add_page_break()

# ===== 1. PRÉSENTATION GÉNÉRALE =====
add_colored_heading(doc, '1. PRÉSENTATION GÉNÉRALE', 1)

doc.add_paragraph(
    'FootCamp Dreams est une application web complète et professionnelle pour la gestion '
    'des réservations d\'un centre de vacances football premium. L\'application permet '
    'aux clients de réserver des séjours, de sélectionner des activités et des services, '
    'tandis que les administrateurs gèrent les réservations, les chambres, les tarifs et '
    'la facturation.'
)

add_colored_heading(doc, '1.1 Objectifs Principaux', 2)
objectives = [
    'Faciliter la réservation de séjours pour les clients',
    'Automatiser la gestion des réservations côté administrateur',
    'Optimiser l\'allocation des chambres',
    'Générer automatiquement les factures et les devis',
    'Gérer les activités et les installations',
    'Authentifier les utilisateurs de manière sécurisée',
    'Envoyer des notifications par email'
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

add_colored_heading(doc, '1.2 Informations du Projet', 2)
info_rows = [
    ['Nom du projet', 'FootCamp Dreams'],
    ['Type', 'Application Web - SPA (Single Page Application)'],
    ['Technologie Frontend', 'HTML5, CSS3, JavaScript (Vanilla + jQuery)'],
    ['Technologie Backend', 'PHP 7.4+'],
    ['Stockage des données', 'JSON (Fichiers plats)'],
    ['Design', 'Responsive - Mobile First'],
    ['Framework CSS', 'Bootstrap 5.3.3'],
    ['Icônes', 'Font Awesome 6.4.0'],
    ['Polices', 'Bebas Neue, DM Sans'],
    ['Version actuelle', '2.0'],
    ['Statut', 'Production Ready']
]
add_table_with_data(doc, ['Caractéristique', 'Détail'], info_rows, 'C0E0DE')

# ===== 2. ARCHITECTURE TECHNIQUE =====
doc.add_page_break()
add_colored_heading(doc, '2. ARCHITECTURE TECHNIQUE', 1)

add_colored_heading(doc, '2.1 Structure du Projet', 2)
doc.add_paragraph(
    'Le projet suit une architecture modulaire et organisée pour une maintenance optimale:'
)

structure = [
    ('📁 mm/', 'Répertoire racine'),
    ('  ├── 📄 index.html', 'Landing page et formulaire de réservation'),
    ('  ├── 📄 client-dashboard.html', 'Tableau de bord client'),
    ('  ├── 📄 admin-dashboard.html', 'Tableau de bord administrateur'),
    ('  ├── 📄 config.php', 'Configuration SMTP'),
    ('  ├── 📁 api/', 'API REST endpoints'),
    ('  │   ├── index.php (Router)', 'Routeur centralisé'),
    ('  │   ├── utils.php', 'Fonctions utilitaires'),
    ('  │   ├── auth/', 'Authentification'),
    ('  │   ├── reservations/', 'Gestion des réservations'),
    ('  │   ├── rooms/', 'Gestion des chambres'),
    ('  │   ├── billing/', 'Facturation et paiements'),
    ('  │   ├── services/', 'Activités et installations'),
    ('  │   ├── admin/', 'Opérations administrateur'),
    ('  │   └── email/', 'Envoi d\'emails'),
    ('  ├── 📁 assets/', 'Ressources statiques'),
    ('  │   ├── style.css', 'Feuille de styles principale'),
    ('  │   ├── admin.js', 'JavaScript admin'),
    ('  │   ├── client.js', 'JavaScript client'),
    ('  │   ├── session-manager.js', 'Gestion des sessions'),
    ('  │   └── *.jpeg', 'Images du centre'),
    ('  ├── 📁 data/', 'Stockage des données JSON'),
    ('  │   ├── users.json', 'Utilisateurs (clients/admins)'),
    ('  │   ├── reservations.json', 'Réservations'),
    ('  │   ├── rooms.json', 'Définition des chambres'),
    ('  │   ├── activities.json', 'Activités disponibles'),
    ('  │   ├── facilities.json', 'Installations (terrains)'),
    ('  │   ├── settings.json', 'Configuration générale'),
    ('  │   └── email_logs.json', 'Journaux d\'emails'),
]

for item in structure:
    p = doc.add_paragraph(item[0], style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

add_colored_heading(doc, '2.2 Stack Technologique', 2)
stack_rows = [
    ['Frontend', 'HTML5, CSS3, JavaScript (ES6+)'],
    ['Framework Frontend', 'Bootstrap 5.3.3'],
    ['JavaScript', 'jQuery 3.5.1, Vanilla JS'],
    ['Icônes', 'Font Awesome 6.4.0'],
    ['Backend', 'PHP 7.4+'],
    ['Base de données', 'JSON (fichiers plats)'],
    ['Stockage', 'Système de fichiers'],
    ['Communication', 'AJAX/Fetch API'],
    ['Email', 'SMTP (Gmail/Custom)'],
    ['Chiffrement mot de passe', 'bcrypt (PHP)'],
    ['Responsive Design', 'Mobile First'],
]
add_table_with_data(doc, ['Composant', 'Technologie'], stack_rows, 'E8F4F8')

# ===== 3. FONCTIONNALITÉS PRINCIPALES =====
doc.add_page_break()
add_colored_heading(doc, '3. FONCTIONNALITÉS PRINCIPALES', 1)

add_colored_heading(doc, '3.1 Pour les Clients', 2)

add_colored_heading(doc, '3.1.1 Réservation', 3)
doc.add_paragraph(
    'Les clients peuvent créer une réservation en remplissant un formulaire avec:'
)
client_res_features = [
    'Informations personnelles (nom, prénom, email, téléphone)',
    'Dates d\'arrivée et de départ',
    'Nombre de personnes',
    'Sélection des types de chambres (simple, double, triple)',
    'Choix des activités (entraînement, match, cours privé, tournoi)',
    'Sélection des installations (stade principal, terrain 7v7, 5v5)',
    'Services additionnels (repas, transport, etc.)',
    'Validation automatique de la saisie'
]
for feature in client_res_features:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.1.2 Authentification', 3)
doc.add_paragraph(
    'Un mot de passe est généré automatiquement lors de la création de la réservation:'
)
auth_features = [
    'Génération de mot de passe 12 caractères (majuscules, minuscules, chiffres, caractères spéciaux)',
    'Chiffrement bcrypt des mots de passe',
    'Affichage modal avec identifiants et boutons de copie',
    'Connexion sécurisée au dashboard client',
    'Récupération de mot de passe oubli'
]
for feature in auth_features:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.1.3 Dashboard Client', 3)
client_dashboard = [
    'Affichage de la réservation en cours',
    'Détails de la réservation (dates, chambres, activités)',
    'Statut de la réservation (en attente, validée, rejetée)',
    'Numéro de chambre assignée (quand validée)',
    'Facture détaillée avec décomposition des coûts',
    'Historique des actions',
    'Modification de profil'
]
for feature in client_dashboard:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.2 Pour les Administrateurs', 2)

add_colored_heading(doc, '3.2.1 Gestion des Réservations', 3)
admin_res_features = [
    'Vue d\'ensemble de toutes les réservations en attente',
    'Validation/Rejet de réservations',
    'Affichage des détails complets de chaque réservation',
    'Calendrier de disponibilité des chambres',
    'Allocation automatique ou manuelle des chambres',
    'Historique des modifications',
    'Filtrage par statut, date, client'
]
for feature in admin_res_features:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.2.2 Gestion des Chambres', 3)
rooms_features = [
    'Trois types de chambres: Simple (1 pers), Double (2 pers), Triple (3 pers)',
    'Affichage du stock total et disponible par type',
    'Tarification par type et par nuit',
    'Visualisation du calendrier des réservations',
    'Assignation intelligente des chambres',
    'Gestion des prix et des capacités'
]
for feature in rooms_features:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.2.3 Gestion Financière', 3)
financial_features = [
    'Génération automatique de factures',
    'Calcul des tarifs (chambres + activités + services)',
    'Gestion des acomptes/dépôts',
    'Application de réductions et promotions',
    'Enregistrement des paiements',
    'Suivi du statut de paiement',
    'Export des factures en PDF'
]
for feature in financial_features:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.2.4 Gestion des Activités et Installations', 3)
activities_features = [
    'Configuration des 4 activités principales',
    'Définition des 3 installations (terrains)',
    'Tarification par activité/installation',
    'Planification des activités par jour',
    'Gestion des participants minimum/maximum',
    'Lien entre activités et installations'
]
for feature in activities_features:
    doc.add_paragraph(feature, style='List Bullet')

add_colored_heading(doc, '3.2.5 Configuration SMTP', 3)
doc.add_paragraph(
    'Interface de configuration pour l\'envoi d\'emails:'
)
smtp_features = [
    'Configuration du serveur SMTP',
    'Authentification Gmail ou autre serveur',
    'Test d\'envoi d\'email',
    'Logs des emails envoyés',
    'Messages automatiques (confirmation, validation, rejet)'
]
for feature in smtp_features:
    doc.add_paragraph(feature, style='List Bullet')

# ===== 4. INTERFACE UTILISATEUR =====
doc.add_page_break()
add_colored_heading(doc, '4. INTERFACE UTILISATEUR', 1)

add_colored_heading(doc, '4.1 Design et Esthétique', 2)

design_info = [
    ['Palette de couleurs', 'Vert (#00E676), Or (#FFD740), Noir (#070A0F)'],
    ['Typographie', 'Bebas Neue (titres), DM Sans (corps)'],
    ['Responsive', 'Mobile-first, Tablet, Desktop'],
    ['Animations', 'Transitions lisses, orbes animées'],
    ['Thème', 'Dark mode moderne et professionnel'],
    ['Framework CSS', 'Bootstrap 5.3.3 + CSS custom'],
]
add_table_with_data(doc, ['Élément', 'Détail'], design_info, 'F0E68C')

add_colored_heading(doc, '4.2 Composants UI Principaux', 2)

components = [
    'Navigation responsive avec menu hamburger',
    'Carrousels d\'images (installations, chambres)',
    'Formulaires avec validation en temps réel',
    'Modals pour affichage d\'identifiants',
    'Tableaux de données avec tri/filtrage',
    'Graphiques (Chart.js) pour statistiques',
    'Badges de statut colorés',
    'Alertes visuelles (succès, erreur, info)',
    'Loading spinners'
]
for comp in components:
    doc.add_paragraph(comp, style='List Bullet')

add_colored_heading(doc, '4.3 Sections de la Landing Page', 2)

landing_sections = [
    ('Hero Section', 'Titre principal "FOOTCAMP DREAMS" avec CTA'),
    ('Feature Cards', '6 cartes présentant les services (6 icônes FA)'),
    ('Galerie Images', 'Carrouels des installations et chambres'),
    ('Statistiques', '4 statistiques clés du centre'),
    ('CTA Section', 'Appels à action avec boutons'),
    ('Formulaire de Réservation', 'Formulaire principal de booking'),
    ('Portails Client/Admin', 'Zones de connexion'),
    ('Footer', 'Adresse, contact, réseaux sociaux'),
]
for section, desc in landing_sections:
    p = doc.add_paragraph(f'{section}: {desc}', style='List Bullet')

# ===== 5. API ET ENDPOINTS =====
doc.add_page_break()
add_colored_heading(doc, '5. API ET ENDPOINTS', 1)

add_colored_heading(doc, '5.1 Architecture API', 2)
doc.add_paragraph(
    'L\'API suit une architecture RESTful avec un routeur centralisé (api/index.php) '
    'qui redirige les appels vers les endpoints appropriés. Les données sont '
    'échangées en format JSON.'
)

add_colored_heading(doc, '5.2 Authentification', 2)
auth_endpoints = [
    ['POST', 'api/login.php', 'Connexion utilisateur'],
    ['POST', 'api/register.php', 'Inscription nouveau client'],
    ['POST', 'api/forgot_password.php', 'Récupération mot de passe'],
    ['POST', 'api/admin_login.php', 'Connexion administrateur'],
    ['POST', 'api/admin_register.php', 'Création admin (admin only)'],
]
add_table_with_data(doc, ['Méthode', 'Endpoint', 'Description'], auth_endpoints, 'D5F4E6')

add_colored_heading(doc, '5.3 Réservations', 2)
reservation_endpoints = [
    ['POST', 'api/reservation_request.php', 'Créer une nouvelle réservation'],
    ['POST', 'api/reserve.php', 'Valider une réservation'],
    ['GET', 'api/get_my_reservations.php', 'Récupérer mes réservations'],
    ['GET', 'api/admin_list_reservations.php', 'Lister toutes les réservations (admin)'],
    ['POST', 'api/admin_validate_reservation.php', 'Valider une réservation (admin)'],
    ['POST', 'api/admin_validate_reservation_with_room.php', 'Valider avec assignation chambre'],
    ['POST', 'api/admin_reject_reservation.php', 'Rejeter une réservation (admin)'],
]
add_table_with_data(doc, ['Méthode', 'Endpoint', 'Description'], reservation_endpoints, 'D5F4E6')

add_colored_heading(doc, '5.4 Chambres', 2)
rooms_endpoints = [
    ['GET', 'api/get_available_rooms.php', 'Lister les chambres disponibles'],
    ['GET', 'api/get_room_calendar.php', 'Calendrier d\'occupation'],
    ['POST', 'api/assign_room_number.php', 'Assigner une chambre'],
    ['POST', 'api/auto_assign_rooms.php', 'Assignation automatique intelligente'],
]
add_table_with_data(doc, ['Méthode', 'Endpoint', 'Description'], rooms_endpoints, 'D5F4E6')

add_colored_heading(doc, '5.5 Facturation', 2)
billing_endpoints = [
    ['GET', 'api/invoice.php', 'Générer facture'],
    ['GET', 'api/get_invoice.php', 'Récupérer facture'],
    ['POST', 'api/record_payment.php', 'Enregistrer paiement'],
    ['GET', 'api/get_payment_status.php', 'Statut de paiement'],
    ['POST', 'api/send_invoice.php', 'Envoyer facture (email)'],
]
add_table_with_data(doc, ['Méthode', 'Endpoint', 'Description'], billing_endpoints, 'D5F4E6')

add_colored_heading(doc, '5.6 Activités et Installations', 2)
services_endpoints = [
    ['GET', 'api/get_activities.php', 'Lister les activités'],
    ['GET', 'api/get_available_facilities.php', 'Lister les installations'],
    ['POST', 'api/add_service.php', 'Ajouter service (admin)'],
]
add_table_with_data(doc, ['Méthode', 'Endpoint', 'Description'], services_endpoints, 'D5F4E6')

# ===== 6. GESTION DES DONNÉES =====
doc.add_page_break()
add_colored_heading(doc, '6. GESTION DES DONNÉES', 1)

add_colored_heading(doc, '6.1 Format de Stockage', 2)
doc.add_paragraph(
    'Tous les données sont stockées en format JSON dans le répertoire /data/. '
    'Cette approche permet une facilité de déploiement et de sauvegarde.'
)

add_colored_heading(doc, '6.2 Fichiers de Données', 2)

data_files = [
    ('users.json', 'Profils clients et administrateurs'),
    ('reservations.json', 'Toutes les réservations avec statuts'),
    ('rooms.json', 'Types et listes des chambres'),
    ('activities.json', 'Catalogue des activités'),
    ('facilities.json', 'Installations (terrains de foot)'),
    ('settings.json', 'Configuration globale de l\'application'),
    ('email_logs.json', 'Historique d\'envoi d\'emails'),
    ('planned_activities.json', 'Activités planifiées par jour'),
]

for filename, description in data_files:
    p = doc.add_paragraph(f'{filename}', style='List Bullet')
    p_child = doc.add_paragraph(description, style='List Bullet 2')

add_colored_heading(doc, '6.3 Structure Exemple: Réservation', 2)
doc.add_paragraph('Exemple de structure JSON d\'une réservation:')

reservation_json = '''{
    "id": 1,
    "nom": "Dupont",
    "prenom": "Jean",
    "email": "jean@example.com",
    "telephone": "0612345678",
    "date_arrivee": "2026-04-09",
    "date_depart": "2026-04-11",
    "nb_personnes": 5,
    "activities": [2, 3],
    "selected_facilities": [1, 2],
    "selected_rooms": {"simple": 0, "double": 1, "triple": 1},
    "status": "validee",
    "deposit": 80,
    "room_numbers": [101, 151],
    "created_at": "2026-04-08 21:18:28"
}'''

code_para = doc.add_paragraph(reservation_json)
code_para.style = 'No Spacing'
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

add_colored_heading(doc, '6.4 Types de Chambres', 2)
room_types = [
    ['Chambre Simple', '1 personne', '50€/nuit', '100 disponibles'],
    ['Chambre Double', '2 personnes', '90€/nuit', '50 disponibles'],
    ['Chambre Triple', '3 personnes', '120€/nuit', '50 disponibles'],
]
add_table_with_data(doc, ['Type', 'Capacité', 'Tarif', 'Stock'], room_types, 'FFE4B5')

add_colored_heading(doc, '6.5 Activités Disponibles', 2)
activities = [
    ['Entraînement collectif', '50€', '22 min', 'Stade Principal'],
    ['Match amical', '75€', '15 min', 'Stade Principal'],
    ['Cours privé', '100€', 'Max 5', 'Tous terrains'],
    ['Tournoi', '200€', '15 min', 'Stade Principal'],
]
add_table_with_data(doc, ['Activité', 'Prix', 'Participants', 'Lieu'], activities, 'D5E8D4')

# ===== 7. SÉCURITÉ ET AUTHENTIFICATION =====
doc.add_page_break()
add_colored_heading(doc, '7. SÉCURITÉ ET AUTHENTIFICATION', 1)

add_colored_heading(doc, '7.1 Authentification', 2)
security_auth = [
    'Hachage bcrypt des mots de passe (cost factor 10)',
    'Génération de mots de passe forts (12 caractères)',
    'Validation des emails (format)',
    'Vérification du numéro de téléphone (10 chiffres)',
    'Comparaison sécurisée des hashs (bcrypt verify)',
    'Prévention du brute force (à implémenter)',
]
for item in security_auth:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, '7.2 Validation des Données', 2)
validation = [
    'Validation côté client (HTML5 form validation)',
    'Validation côté serveur (PHP)',
    'Vérification des dates (arrivée < départ)',
    'Contrôle des formats (email, téléphone)',
    'Vérification des quantités (chambres, personnes)',
    'Validation des montants financiers',
]
for item in validation:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, '7.3 Recommandations de Sécurité', 2)
recommendations = [
    'HTTPS obligatoire en production',
    'Implémenter un système de taux limite (rate limiting)',
    'Ajouter CSRF tokens dans les formulaires',
    'Valider et nettoyer toutes les entrées (input sanitization)',
    'Implémenter des sessions PHP sécurisées',
    'Logger toutes les actions sensibles',
    'Sauvegarder régulièrement les données JSON',
    'Implémenter une 2FA pour les administrateurs',
    'Audit log des modifications d\'administrateur',
]
for item in recommendations:
    doc.add_paragraph(item, style='List Bullet')

# ===== 8. INSTALLATION ET DÉPLOIEMENT =====
doc.add_page_break()
add_colored_heading(doc, '8. INSTALLATION ET DÉPLOIEMENT', 1)

add_colored_heading(doc, '8.1 Prérequis', 2)
requirements = [
    'PHP 7.4 ou supérieur',
    'Extension PHP JSON (incluse par défaut)',
    'Accès en écriture aux répertoires /data/',
    'Navigateur moderne (Chrome, Firefox, Safari, Edge 2020+)',
    'Connexion Internet (CDNs Bootstrap, Font Awesome, Google Fonts)',
]
for req in requirements:
    doc.add_paragraph(req, style='List Bullet')

add_colored_heading(doc, '8.2 Installation Locale', 2)
install_steps = [
    'Télécharger les fichiers du projet',
    'Placer dans un répertoire accessible',
    'Vérifier les permissions de lecture/écriture sur /data/',
    'Lancer un serveur HTTP local (Python: python -m http.server 8000)',
    'Ouvrir http://localhost:8000 dans le navigateur',
]
for i, step in enumerate(install_steps, 1):
    doc.add_paragraph(f'{step}', style='List Number')

add_colored_heading(doc, '8.3 Déploiement en Production', 2)
deployment = [
    'Uploader les fichiers sur le serveur web',
    'Configurer les permissions (755 pour dossiers, 644 pour fichiers)',
    'Configurer PHP 7.4+',
    'Mettre à jour config.php avec les paramètres SMTP',
    'Tester l\'application complet (login, réservation, admin)',
    'Mettre en place un plan de sauvegarde /data/',
    'Configurer HTTPS avec SSL/TLS',
]
for item in deployment:
    doc.add_paragraph(item, style='List Bullet')

add_colored_heading(doc, '8.4 Configuration SMTP', 2)
doc.add_paragraph(
    'Éditer le fichier config.php pour configurer l\'envoi d\'emails:'
)

config_example = '''define('SMTP_HOST', 'smtp.gmail.com');
define('SMTP_PORT', 587);
define('SMTP_SECURE', 'tls');
define('SMTP_USER', 'your-email@gmail.com');
define('SMTP_PASSWORD', 'your-app-password');
define('SMTP_FROM_EMAIL', 'noreply@footcamp.com');
define('SMTP_FROM_NAME', 'FootCamp Dreams');'''

config_para = doc.add_paragraph(config_example)
config_para.style = 'No Spacing'
for run in config_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

# ===== PAGE FINALE =====
doc.add_page_break()
add_colored_heading(doc, 'CONCLUSION', 1)

doc.add_paragraph(
    'FootCamp Dreams est une application web mature et fonctionnelle pour la gestion '
    'complète des réservations d\'un centre de vacances football. L\'application offre:'
)

conclusion_points = [
    'Une interface moderne et responsive adaptée aux mobiles',
    'Une API RESTful bien structurée et documentée',
    'Un système d\'authentification sécurisé',
    'Une gestion complète du cycle de vie des réservations',
    'Une gestion intelligente des ressources (chambres, activités)',
    'Un système de facturation automatisé',
    'Une architecture modulaire et facilement maintenable',
]

for point in conclusion_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'L\'application est prête pour la production et peut être déployée sur '
    'un serveur web standard avec PHP 7.4+.'
)

# Infos finales
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run(f'Rapport généré le {datetime.now().strftime("%d/%m/%Y à %H:%M:%S")}\nFootCamp Dreams - Version 2.0')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(138, 155, 176)

# ===== FLUX UTILISATEUR =====
doc.add_page_break()
add_colored_heading(doc, '9. FLUX UTILISATEUR', 1)

add_colored_heading(doc, '9.1 Flux Client', 2)
client_flow = [
    ('Visite le site', 'index.html - Landing page avec présentation'),
    ('Scroll vers réservation', 'Voit le formulaire de réservation'),
    ('Remplit le formulaire', 'Infos personnelles, dates, activités, services'),
    ('Soumission', 'Clic sur "Soumettre ma réservation"'),
    ('Génération mot de passe', 'Mot de passe aléatoire généré côté client'),
    ('Modal identifiants', 'Affichage avec email, password, numéro réserv.'),
    ('Copie identifiants', 'Possibilité de copier dans le presse-papiers'),
    ('Création utilisateur', 'Stockage dans users.json avec bcrypt'),
    ('Accès Dashboard', 'Redirection vers client-dashboard.html'),
    ('Suivi réservation', 'Voit statut, chambre assignée, facture')
]
for action, detail in client_flow:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(action + ': ').bold = True
    p.add_run(detail)

add_colored_heading(doc, '9.2 Flux Administrateur', 2)
admin_flow = [
    ('Connexion', 'Accès au admin-dashboard.html avec login'),
    ('Vue des réservations', 'Liste toutes les réservations en attente'),
    ('Validation', 'Vérification des détails et validation'),
    ('Assignation chambre', 'Allocation automatique ou manuelle'),
    ('Notification client', 'Email envoyé au client (configurable)'),
    ('Accès client', 'Client peut voir sa chambre et sa facture'),
    ('Paiement', 'Client enregistre son paiement'),
    ('Confirmationfinale', 'Réservation marquée comme complète')
]
for action, detail in admin_flow:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(action + ': ').bold = True
    p.add_run(detail)

# ===== EXEMPLES DE DONNÉES =====
doc.add_page_break()
add_colored_heading(doc, '10. EXEMPLES DE DONNÉES', 1)

add_colored_heading(doc, '10.1 Structure d\'un Utilisateur', 2)
doc.add_paragraph('Exemple utilisateur dans users.json:')
user_example = '''{
    "id": 1,
    "nom": "Dupont",
    "prenom": "Jean",
    "email": "jean@example.com",
    "password": "$2y$10$...",
    "telephone": "0612345678",
    "role": "client"
}'''
code_para = doc.add_paragraph(user_example)
code_para.style = 'No Spacing'
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

add_colored_heading(doc, '10.2 Détail d\'une Facture', 2)
doc.add_paragraph('Calcul automatique des coûts:')
invoice_data = [
    ['Chambre Double (3 nuits)', '90€ × 3', '270€'],
    ['Entraînement collectif', '50€ × 3 jours × 5 pers', '750€'],
    ['Match amical', '75€ × 1 jour × 5 pers', '375€'],
    ['Repas complet', '30€ × 3 jours × 5 pers', '450€'],
    ['Sous-total', '', '1845€'],
    ['Réduction 10%', '', '-184€'],
    ['Acompte versé', '', '-80€'],
    ['Solde dû', '', '1581€'],
]
add_table_with_data(doc, ['Description', 'Calcul', 'Montant'], invoice_data, 'FFF8DC')

# ===== RECOMMANDATIONS =====
doc.add_page_break()
add_colored_heading(doc, '11. RECOMMANDATIONS ET AMÉLIORATIONS FUTURES', 1)

add_colored_heading(doc, '11.1 Sécurité - Priorité Haute', 2)
security_rec = [
    'Forcer HTTPS/TLS en production',
    'Implémenter CSRF tokens dans tous les formulaires',
    'Ajouter rate limiting pour prévenir les attaques brute force',
    'Implémenter 2FA pour les administrateurs',
    'Chiffrer les données sensibles (carte bancaire si applicable)',
    'Audit logging de toutes les opérations administrateur',
    'Implémenter des sessions avec timeouts',
]
for rec in security_rec:
    doc.add_paragraph(rec, style='List Bullet')

add_colored_heading(doc, '11.2 Performance - Priorité Moyenne', 2)
perf_rec = [
    'Implémenter un système de cache (Redis/Memcached)',
    'Compresser les réponses API (gzip)',
    'Minifier CSS et JavaScript',
    'Implémenter lazy loading pour les images',
    'Optimiser les images (WebP)',
    'Ajouter Service Workers pour offline support',
]
for rec in perf_rec:
    doc.add_paragraph(rec, style='List Bullet')

add_colored_heading(doc, '11.3 Fonctionnalités - Priorité Moyenne', 2)
feature_rec = [
    'Implémenter un vrai système de base de données (MySQL/PostgreSQL)',
    'Ajouter un système de paiement en ligne (Stripe, PayPal)',
    'Implémenter des notifications SMS',
    'Ajouter un système de avis/commentaires',
    'Créer une app mobile (React Native, Flutter)',
    'Implémenter un CRM pour le suivi clients',
    'Ajouter un système de fidélité/points',
]
for rec in feature_rec:
    doc.add_paragraph(rec, style='List Bullet')

add_colored_heading(doc, '11.4 DevOps - Priorité Basse', 2)
devops_rec = [
    'Mettre en place CI/CD (GitHub Actions, Jenkins)',
    'Conteneuriser avec Docker',
    'Implémenter des tests automatisés (Jest, PHPUnit)',
    'Ajouter un monitoring (Prometheus, Grafana)',
    'Implémenter des backups automatiques',
    'Mettre en place un système de versioning',
]
for rec in devops_rec:
    doc.add_paragraph(rec, style='List Bullet')

# ===== CONCLUSION =====
doc.add_page_break()
add_colored_heading(doc, '12. CONCLUSION', 1)

doc.add_paragraph(
    'FootCamp Dreams est une application web robuste et bien structurée pour '
    'la gestion des réservations. Voici les points forts:'
)

strengths = [
    'Architecture modulaire et maintenable',
    'Interface utilisateur moderne et responsive',
    'API RESTful bien organisée',
    'Système d\'authentification sécurisé',
    'Gestion complète du cycle de vie des réservations',
    'Allocation intelligente des ressources',
    'Système de facturation automatisé',
    'Documentation complète du projet',
]
for strength in strengths:
    doc.add_paragraph(strength, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph(
    'L\'application est prête pour être déployée en production sur un serveur '
    'PHP 7.4+ et peut gérer efficacement les opérations d\'un centre de '
    'vacances football.'
)

# INFORMATIONS DE CONTACT
doc.add_paragraph()
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run(
    f'\n\nDocument généré le {datetime.now().strftime("%d/%m/%Y à %H:%M:%S")}\n'
    'FootCamp Dreams - Application de Gestion de Réservations\n'
    'Version 2.0 - Production Ready'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(138, 155, 176)

# Sauvegarder le document
output_path = 'D:\\parisss\\WEB\\mm\\web.docx'
doc.save(output_path)
print('✓ Rapport généré avec succès: web.docx')
print(f'✓ Fichier sauvegardé dans: {output_path}')
print('✓ Le rapport couvre 12 sections principales avec tableaux et listes détaillées')

